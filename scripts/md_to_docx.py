# -*- coding: utf-8 -*-
"""Build the competition plan DOCX from docs/项目计划书.md."""

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont

ROOT = Path(r"C:\Users\34036\Desktop\MultiModel-Safety-Hazard-Agent")
MD_PATH = ROOT / "docs" / "项目计划书.md"
DOCX_PATH = ROOT / "多模态基层安全隐患智能研判与处置辅助系统项目计划书.docx"
DIAGRAM_PATH = ROOT / "docs" / "assets" / "architecture.png"

EAST_BODY = "宋体"
EAST_HEAD = "微软雅黑"
ASCII_BODY = "Times New Roman"
ASCII_HEAD = "Microsoft YaHei"
TABLE_WIDTH_CM = 16.2


def set_run(run, size=11, bold=False, color="000000", east=EAST_BODY, ascii_font=ASCII_BODY):
    run.font.name = ascii_font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)
    rfonts.set(qn("w:eastAsia"), east)


def add_text_runs(p, text, size=11, bold=False, color="000000", east=EAST_BODY, ascii_font=ASCII_BODY):
    parts = str(text).split("**")
    for index, part in enumerate(parts):
        if not part:
            continue
        run = p.add_run(part)
        set_run(run, size=size, bold=(index % 2 == 1) or bold, color=color, east=east, ascii_font=ascii_font)


def add_para(doc, text, size=11, bold=False, after=5, indent=None, align=None):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = 1.35
    pf.space_after = Pt(after)
    if indent is not None:
        pf.left_indent = Cm(indent)
    if align is None:
        align = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.alignment = align
    add_text_runs(p, text, size=size, bold=bold)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Cm(0.7)
    pf.first_line_indent = Cm(-0.35)
    pf.line_spacing = 1.35
    pf.space_after = Pt(3)
    r = p.add_run("•  ")
    set_run(r, size=11, east=EAST_BODY)
    add_text_runs(p, text, size=11)
    return p


def add_num(doc, n, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Cm(0.7)
    pf.first_line_indent = Cm(-0.5)
    pf.line_spacing = 1.35
    pf.space_after = Pt(3)
    r = p.add_run(f"{n}.  ")
    set_run(r, size=11, east=EAST_BODY)
    add_text_runs(p, text, size=11)
    return p


def add_heading(doc, text, level):
    style = "Heading 1" if level == 1 else "Heading 2"
    p = doc.add_paragraph(style=style)
    pf = p.paragraph_format
    pf.space_before = Pt(14 if level == 1 else 10)
    pf.space_after = Pt(6 if level == 1 else 4)
    pf.keep_with_next = True
    pf.line_spacing = 1.2
    add_text_runs(p, text, size=15 if level == 1 else 12.5, bold=True, east=EAST_HEAD, ascii_font=ASCII_HEAD)
    return p


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_table_borders(table, color="D9D9D9"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    tbl_pr.append(borders)


def set_cell_margins(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    mar = OxmlElement("w:tblCellMar")
    for edge, width in (("top", "60"), ("left", "110"), ("bottom", "60"), ("right", "110")):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:w"), width)
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tbl_pr.append(mar)


def set_fixed_layout(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    el = OxmlElement("w:tblHeader")
    el.set(qn("w:val"), "true")
    tr_pr.append(el)


def compute_widths(rows):
    ncols = max(len(r) for r in rows)
    lens = []
    for ci in range(ncols):
        mx = 1
        for r in rows:
            if ci < len(r):
                mx = max(mx, len(str(r[ci])) + 2)
        lens.append(mx)
    total = sum(lens)
    widths = [max(1.6, min(7.2, TABLE_WIDTH_CM * l / total)) for l in lens]
    factor = TABLE_WIDTH_CM / sum(widths)
    return [round(w * factor, 2) for w in widths]


def col_align(rows, ci):
    for r in rows:
        if ci < len(r) and len(str(r[ci])) > 9:
            return "left"
    return "center"


def add_cell_text(cell, text, size=9.5, bold=False, color="000000", align="left"):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    first = True
    for line in str(text).split("\n"):
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        pf = p.paragraph_format
        pf.line_spacing = 1.15
        pf.space_before = Pt(1)
        pf.space_after = Pt(1)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if align == "center" else WD_ALIGN_PARAGRAPH.LEFT
        add_text_runs(p, line, size=size, bold=bold, color=color)


def add_table(doc, rows):
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = compute_widths(rows)
    aligns = [col_align(rows, ci) for ci in range(ncols)]
    set_table_borders(table)
    set_cell_margins(table)
    set_fixed_layout(table)
    headers = rows[0]
    for ci, header in enumerate(headers):
        cell = table.rows[0].cells[ci]
        cell.width = Cm(widths[ci])
        shade_cell(cell, "1F4E79")
        add_cell_text(cell, header, size=9.5, bold=True, color="FFFFFF", align="center")
    repeat_header(table.rows[0])
    for ri, row in enumerate(rows[1:], start=1):
        for ci in range(ncols):
            cell = table.rows[ri].cells[ci]
            cell.width = Cm(widths[ci])
            if ri % 2 == 0:
                shade_cell(cell, "F2F6FA")
            text = row[ci] if ci < len(row) else ""
            add_cell_text(cell, text, size=9.5, align=aligns[ci])
    return table


def add_code_block(doc, text):
    lines = text.strip("\n").split("\n")
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, color="CFCFCF")
    set_cell_margins(table)
    set_fixed_layout(table)
    cell = table.rows[0].cells[0]
    cell.width = Cm(TABLE_WIDTH_CM)
    shade_cell(cell, "F5F5F5")
    first = True
    for line in lines:
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        pf = p.paragraph_format
        pf.line_spacing = 1.12
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        r = p.add_run(line if line else " ")
        set_run(r, size=8.5, east=EAST_HEAD, ascii_font="Consolas")
    add_para(doc, "", size=6, after=2)


def add_figure(doc, path, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run()
    run.add_picture(str(path), width=Cm(15.8))
    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraph_format.space_after = Pt(6)
    add_text_runs(c, caption, size=9.5, bold=True)


def collect_table(lines, start):
    rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        raw = lines[i].strip()
        cells = [c.strip() for c in raw.strip("|").split("|")]
        rows.append(cells)
        i += 1
    if len(rows) >= 2 and all(set(c) <= set("-: ") for c in rows[1]):
        rows.pop(1)
    return rows, i


def parse_md(path):
    lines = path.read_text(encoding="utf-8").splitlines()
    blocks = []
    cover = True
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line.strip():
            i += 1
            continue
        if line.strip() == "---":
            if cover:
                blocks.append(("pagebreak",))
                cover = False
            i += 1
            continue
        if cover:
            if line.startswith("# "):
                blocks.append(("title", line[2:].strip()))
            elif line.startswith("**") and line.endswith("**"):
                blocks.append(("subtitle", line.strip("*")))
            elif line.startswith("|"):
                rows, i = collect_table(lines, i)
                blocks.append(("table", rows))
                continue
            else:
                blocks.append(("cover_line", line.strip()))
            i += 1
            continue
        if line.startswith("# "):
            blocks.append(("h1", line[2:].strip()))
        elif line.startswith("## "):
            blocks.append(("h2", line[3:].strip()))
        elif line.startswith("```"):
            code = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                code.append(lines[i].rstrip("\n"))
                i += 1
            blocks.append(("code", "\n".join(code)))
        elif line.startswith("|"):
            rows, i = collect_table(lines, i)
            blocks.append(("table", rows))
            continue
        else:
            m = re.match(r"^!\[(.+)\]\((.+)\)$", line)
            if m:
                blocks.append(("figure", m.group(2), m.group(1)))
            elif line.startswith("- "):
                blocks.append(("bullet", line[2:].strip()))
            else:
                m = re.match(r"^(\d+)\.\s+(.*)$", line)
                if m:
                    blocks.append(("num", int(m.group(1)), m.group(2).strip()))
                else:
                    blocks.append(("para", line.strip()))
        i += 1
    return blocks


def prepare_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)
    normal = doc.styles["Normal"]
    normal.font.name = ASCII_BODY
    normal.font.size = Pt(11)
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), ASCII_BODY)
    rfonts.set(qn("w:hAnsi"), ASCII_BODY)
    rfonts.set(qn("w:eastAsia"), EAST_BODY)
    normal.paragraph_format.line_spacing = 1.35
    normal.paragraph_format.space_after = Pt(5)
    return doc


def add_cover(doc, blocks):
    for block in blocks:
        kind = block[0]
        if kind == "title":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(54)
            p.paragraph_format.space_after = Pt(12)
            p.paragraph_format.line_spacing = 1.25
            add_text_runs(p, block[1], size=21, bold=True, east=EAST_HEAD, ascii_font=ASCII_HEAD)
        elif kind == "subtitle":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(8)
            add_text_runs(p, block[1], size=12, east=EAST_HEAD, ascii_font=ASCII_HEAD)
        elif kind == "cover_line":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(14)
            add_text_runs(p, block[1], size=10.5)
        elif kind == "table":
            label = doc.add_paragraph()
            label.alignment = WD_ALIGN_PARAGRAPH.CENTER
            label.paragraph_format.space_before = Pt(16)
            label.paragraph_format.space_after = Pt(6)
            add_text_runs(label, "项目概览", size=13, bold=True, east=EAST_HEAD, ascii_font=ASCII_HEAD)
            add_table(doc, block[1])
        elif kind == "pagebreak":
            doc.add_page_break()
            return


def build_docx(blocks, out_path, diagram_path):
    doc = prepare_document()
    add_cover(doc, blocks[: blocks.index(("pagebreak",)) + 1])
    seen_figure = False
    for block in blocks[blocks.index(("pagebreak",)) + 1 :]:
        kind = block[0]
        if kind == "h1":
            add_heading(doc, block[1], 1)
        elif kind == "h2":
            add_heading(doc, block[1], 2)
        elif kind == "para":
            add_para(doc, block[1])
        elif kind == "bullet":
            add_bullet(doc, block[1])
        elif kind == "num":
            add_num(doc, block[1], block[2])
        elif kind == "table":
            add_table(doc, block[1])
            add_para(doc, "", size=6, after=2)
        elif kind == "code":
            add_code_block(doc, block[1])
        elif kind == "figure":
            caption = block[2] if len(block) > 2 else ""
            if not seen_figure and diagram_path.exists() and diagram_path.suffix.lower() == ".png":
                add_figure(doc, diagram_path, caption)
                seen_figure = True
    doc.save(out_path)


def load_font(size, bold=False):
    candidates = (
        [r"C:\Windows\Fonts\msyhbd.ttc"] if bold else [r"C:\Windows\Fonts\msyh.ttc"]
    ) + [r"C:\Windows\Fonts\simhei.ttf", r"C:\Windows\Fonts\msyh.ttc", r"C:\Windows\Fonts\simsun.ttc"]
    for candidate in candidates:
        if Path(candidate).exists():
            try:
                return ImageFont.truetype(candidate, size)
            except Exception:
                continue
    return ImageFont.load_default()


def make_diagram(out_path):
    layers = [
        ("交互层", "Vue 3 + Element Plus Web 界面", "照片与描述上传 · 证据链展示 · 历史记录", "#EAF2FB"),
        ("服务与接口层", "FastAPI · REST API", "任务管理 · 文件上传 · 权限 · 导出", "#E7F4EA"),
        ("智能体编排层", "LangGraph Agent 工作流", "解析 → 追问 → 检索 → 研判 → 生成 → 复核", "#FFF4E3"),
        ("知识检索层", "RAG：BGE-M3 + FAISS", "文档切片 · 混合检索 · 重排序 · 条款溯源", "#F2ECF9"),
        ("模型推理层", "Qwen2.5-VL · GLM-4V-Flash", "图像理解 · 结构化生成 · 规则引擎双通道", "#FDECEB"),
        ("数据层", "SQLite / PostgreSQL · 本地存储 / MinIO", "研判记录 · 知识元数据 · 图片与导出文件", "#E7F3F3"),
    ]
    width, height = 2400, 1500
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    f_title = load_font(62, bold=True)
    f_tech = load_font(46)
    f_desc = load_font(40)
    box_w = 1560
    box_h = 218
    gap = 26
    x0 = (width - box_w) // 2
    y0 = 48
    for i, (name, tech, desc, fill) in enumerate(layers):
        y = y0 + i * (box_h + gap)
        draw.rounded_rectangle([x0, y, x0 + box_w, y + box_h], radius=24, fill=fill, outline="#5B7A9D", width=4)
        draw.text((x0 + 52, y + 28), name, font=f_title, fill="#17324D")
        draw.text((x0 + 52, y + 112), tech, font=f_tech, fill="#333333")
        draw.text((x0 + 52, y + 168), desc, font=f_desc, fill="#555555")
        if i < len(layers) - 1:
            ax = x0 + box_w // 2
            ay = y + box_h
            by = y0 + (i + 1) * (box_h + gap)
            draw.line([ax, ay + 3, ax, by - 14], fill="#5B7A9D", width=6)
            draw.polygon([(ax - 18, by - 30), (ax + 18, by - 30), (ax, by - 7)], fill="#5B7A9D")
    out_path.parent.mkdir(parents=True, exist_ok=True)
    img.save(out_path)


def main():
    make_diagram(DIAGRAM_PATH)
    blocks = parse_md(MD_PATH)
    build_docx(blocks, DOCX_PATH, DIAGRAM_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
